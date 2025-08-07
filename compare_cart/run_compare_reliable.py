import sys, os
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
sys.path.append(BASE_DIR)

from flask import current_app  # ✅ только этот импорт!
from docx import Document
from models import db, Part
from compare_cart.reliable_scraper import get_cart_items

def get_reliable_items():
    return get_cart_items()

def check_cart_items(items):
    results = []
    with current_app.app_context():  # ✅ исправлено
        for item in items:
            part = Part.query.filter_by(part_number=item["part_number"]).filter(Part.quantity > 0).first()
            if part:
                results.append({
                    "part_number": item["part_number"],
                    "qty_in_cart": item["qty"],
                    "in_stock": True,
                    "qty_in_stock": part.quantity,
                    "location": part.location,
                    "name": part.name
                })
            else:
                results.append({
                    "part_number": item["part_number"],
                    "qty_in_cart": item["qty"],
                    "in_stock": False,
                    "qty_in_stock": 0,
                    "location": "-",
                    "name": "-"
                })
    return results

def export_to_docx(results, filename="reliable_inventory_report.docx"):
    if os.path.exists(filename):
        os.remove(filename)

    document = Document()
    document.add_heading('Reliable Cart vs Inventory Report', 0)

    table = document.add_table(rows=1, cols=6)
    table.style = 'Light List'
    hdr = table.rows[0].cells
    hdr[0].text = "Part Number"
    hdr[1].text = "Qty in Cart"
    hdr[2].text = "In Stock?"
    hdr[3].text = "Qty in Stock"
    hdr[4].text = "Location"
    hdr[5].text = "Name"

    for r in results:
        row = table.add_row().cells
        row[0].text = r["part_number"]
        row[1].text = str(r["qty_in_cart"])
        row[2].text = "Yes" if r["in_stock"] else "No"
        row[3].text = str(r["qty_in_stock"])
        row[4].text = r["location"]
        row[5].text = r["name"]

    document.save(filename)
    print(f"\n✅ Report saved: {filename}")

if __name__ == "__main__":
    print("🔍 Сбор корзины с Reliable...")
    cart_items = get_cart_items()

    print("📦 Сравнение с базой...")
    result = check_cart_items(cart_items)

    print("📝 Экспорт в Word...")
    export_to_docx(result)
